"""
Resume parser — extracts plain text from PDF / DOCX / TXT / image files.

Design notes:
  - We try multiple PDF backends (pdfplumber first, PyPDF2 fallback) because
    real-world resumes break parsers in creative ways (two-column layouts,
    embedded fonts, scanned pages).
  - Image OCR uses pytesseract, which requires the Tesseract binary. If it's
    not installed we fail gracefully with a clear error message.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a resume file.

    Args:
        file_bytes: Raw file content.
        filename:   Original filename (used to pick parser by extension).

    Returns:
        Cleaned plain text.

    Raises:
        ValueError: If file extension is unsupported.
        RuntimeError: If the chosen parser fails (e.g. corrupt PDF).
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    logger.info("Parsing resume: %s (ext=%s, size=%d bytes)",
                filename, ext, len(file_bytes))

    if ext == "pdf":
        text = _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        text = _parse_docx(file_bytes)
    elif ext in ("txt", "md"):
        text = file_bytes.decode("utf-8", errors="ignore")
    elif ext in ("png", "jpg", "jpeg", "tif", "tiff", "bmp"):
        text = _parse_image(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: .{ext}. "
            "Supported: PDF, DOCX, TXT, PNG, JPG."
        )

    cleaned = _clean_text(text)
    if len(cleaned) < 50:
        raise RuntimeError(
            "Could not extract meaningful text from resume. "
            "If it's a scanned PDF, try uploading as an image instead."
        )
    return cleaned


# ---------------------------------------------------------------------------
# Format-specific parsers
# ---------------------------------------------------------------------------
def _parse_pdf(file_bytes: bytes) -> str:
    """Try pdfplumber first (better with layout), fall back to PyPDF2."""
    # Attempt 1: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(pages)
            if text.strip():
                return text
    except Exception as e:
        logger.warning("pdfplumber failed: %s", e)

    # Attempt 2: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as e:
        logger.error("PyPDF2 also failed: %s", e)
        raise RuntimeError(
            "Unable to read this PDF. It may be scanned or password-protected."
        ) from e


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is not installed.") from e

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Tables often contain skills/experience -- include them too.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def _parse_image(file_bytes: bytes) -> str:
    """OCR via Tesseract. Tesseract binary must be installed on the system."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise RuntimeError(
            "OCR libraries not installed. Run: pip install pytesseract pillow"
        ) from e

    try:
        image = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError as e:
        raise RuntimeError(
            "Tesseract OCR engine is not installed on this system. "
            "Install it from https://github.com/tesseract-ocr/tesseract "
            "or upload your resume as PDF/DOCX instead."
        ) from e


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------
def _clean_text(text: str) -> str:
    """Normalize whitespace and strip control characters."""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Strip non-printable characters but keep newlines and tabs
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
    # Collapse repeated whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Resume metadata extraction (lightweight; richer NLP lives in nlp module)
# ---------------------------------------------------------------------------
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{8,}\d)")
URL_RE = re.compile(r"https?://[^\s)]+")
YEARS_EXP_RE = re.compile(
    r"(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)?",
    re.IGNORECASE,
)


def extract_metadata(text: str) -> dict:
    """Pull structured bits (email, phone, links, years of experience)."""
    emails = list(set(EMAIL_RE.findall(text)))
    phones = list(set(PHONE_RE.findall(text)))
    urls = list(set(URL_RE.findall(text)))

    # Years of experience: take the maximum number we see
    years_matches = [int(m) for m in YEARS_EXP_RE.findall(text) if int(m) < 50]
    years_exp = max(years_matches) if years_matches else None

    return {
        "emails": emails,
        "phones": phones,
        "urls": urls,
        "years_experience": years_exp,
    }
